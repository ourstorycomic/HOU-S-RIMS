# routes/document.py
from flask import Blueprint, request, send_file
import os
import io
import base64
import pandas as pd
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

document_bp = Blueprint('document', __name__)

def set_font(run, size=13, bold=False, italic=False, font_name='Times New Roman', color=None):
    """Thiết lập font chữ chuẩn NCKH"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level, numbering=""):
    """Thêm tiêu đề có đánh số và style chuẩn"""
    h = doc.add_heading('', level)
    run = h.add_run(f"{numbering} {text}".strip())
    run.font.color.rgb = RGBColor(0, 51, 102)
    set_font(run, size=16 if level==1 else 14, bold=True)
    return h

def set_cell_background(cell, fill):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_ai_content(text):
    """Xóa bỏ các tiêu đề chương và tag markers lặp lại từ AI"""
    if not text: return ""
    # Xóa các dòng bắt đầu bằng Chương, Chapter, I, II, III, IV, V...
    text = re.sub(r'^(Chương|Chapter|I|II|III|IV|V)\s*.*[:.-]', '', text, flags=re.IGNORECASE | re.MULTILINE)
    # Xóa các tag Marker gán cho AI
    text = re.sub(r'\[(MÔ TẢ CHI TIẾT|PHƯƠNG PHÁP|CHUYÊN SÂU|KÌ VỌNG & KẾT LUẬN|MÔ TẢ|KẾT LUẬN)\]:?', '', text)
    return text.replace('*', '').replace('#', '').strip()

def add_ai_paragraphs(doc, text):
    """Thêm các đoạn văn từ AI và căn lề justify, thụt đầu dòng chuyên nghiệp"""
    cleaned = clean_ai_content(text)
    for section in cleaned.split('\n\n'):
        if section.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Thụt đầu dòng 1.27cm chuẩn Word
            p.paragraph_format.first_line_indent = Cm(1.27)
            set_font(p.add_run(section.strip()))

def get_stats_table(df, col):
    try:
        val_counts = df[col].value_counts()
        total = len(df)
        stats = []
        for idx, val in val_counts.items():
            pct = (val / total) * 100
            stats.append([str(idx), str(val), f"{pct:.2f}%"])
        return stats
    except: return []

@document_bp.route("/api/generate_doc", methods=["POST"])
def generate_doc():
    try:
        data = request.get_json()
        image_data = data.get('image')
        col_names = [c.strip() for c in data.get('col_name', '').split(',') if c.strip()]
        dataset_id = data.get('dataset_id')
        total_rows = data.get('total_rows', 0)
        total_cols = data.get('total_cols', 0)
        
        # Phân đoạn phân tích
        analysis_ii = data.get('analysis_ii', '')
        analysis_meth = data.get('analysis_meth', '')
        analysis_iv = data.get('analysis_iv', '')
        analysis_v = data.get('analysis_v', '')

        stats_type = data.get('stats_type')
        stats_data = data.get('stats_data')
        
        df = None
        if dataset_id:
            path = os.path.join(os.getcwd(), 'uploads', dataset_id)
            if not os.path.exists(path):
                for ext in ['.xlsx', '.csv', '.sav']:
                    if os.path.exists(path + ext): path += ext; break
            if os.path.exists(path):
                if path.endswith('.csv'): df = pd.read_csv(path)
                elif path.endswith('.sav'): 
                    import pyreadstat
                    df, _ = pyreadstat.read_sav(path)
                else: df = pd.read_excel(path, engine='openpyxl')

        doc = Document()
        for s in doc.sections:
            s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(2), Cm(2), Cm(3), Cm(2)

        # 1. TRANG BÌA
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC MỞ HÀ NỘI"), size=14, bold=True)
        doc.add_paragraph("\n" * 4)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run("BÁO CÁO KẾT QUẢ NGHIÊN CỨU KHOA HỌC"), size=18, bold=True)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(f"ĐỀ TÀI: PHÂN TÍCH DỮ LIỆU ĐỊNH LƯỢNG TRÊN HỆ THỐNG HOU-S-RIMS"), size=16, bold=True, color=RGBColor(0, 51, 102))
        doc.add_paragraph("\n" * 6)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(p.add_run(f"Người thực hiện: Sinh viên HOU\nĐơn vị: Khoa Công nghệ thông tin\nMã định danh: {dataset_id[:8] if dataset_id else 'N/A'}"), size=12, italic=True)
        doc.add_paragraph("\n" * 8)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run("Hà Nội, Năm 2026"), size=12, bold=True)
        doc.add_page_break()

        # 2. MỤC LỤC
        add_heading(doc, "MỤC LỤC", 1)
        p = doc.add_paragraph(); run = p.add_run("(Chuột phải -> Update Field để cập nhật số trang)"); set_font(run, size=11, italic=True, color=RGBColor(255, 0, 0))
        p = doc.add_paragraph(); run = p.add_run()
        for tag in ['begin', 'separate', 'end']:
            f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), tag)
            if tag == 'separate':
                instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
                run._element.append(instr)
            run._element.append(f)
        doc.add_page_break()

        # I.
        add_heading(doc, 'TRỰC QUAN HÓA DỮ LIỆU', 1, "I.")
        if image_data:
            try:
                if "," in image_data: image_data = image_data.split(",")[1]
                doc.add_picture(io.BytesIO(base64.b64decode(image_data)), width=Inches(5.5))
                c = doc.add_paragraph('Hình 1. Biểu đồ mô tả phân bố và xu hướng dữ liệu'); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except: doc.add_paragraph("[Lỗi kết xuất biểu đồ]")

        # II.
        add_heading(doc, 'THỐNG KÊ MÔ TẢ & TẦN SUẤT', 1, "II.")
        if df is not None and col_names:
            for idx, col in enumerate(col_names):
                if col in df.columns:
                    doc.add_heading(f'2.{idx+1}. Thống kê biến: {col}', level=2)
                    stats = get_stats_table(df, col)
                    if stats:
                        t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
                        hdr = t.rows[0].cells; hdr[0].text, hdr[1].text, hdr[2].text = 'Phân loại', 'Tần số (n)', 'Tỷ lệ (%)'
                        for cell in hdr: set_cell_background(cell, "E7E6E6")
                        for item in stats:
                            r = t.add_row().cells; r[0].text, r[1].text, r[2].text = item
        add_ai_paragraphs(doc, analysis_ii)

        # III. NÂNG CẤP
        add_heading(doc, 'PHƯƠNG PHÁP & TỔNG QUAN DỮ LIỆU', 1, "III.")
        doc.add_heading('3.1. Tổng quan cấu trúc bộ dữ liệu', level=2)
        table = doc.add_table(rows=4, cols=2); table.style = 'Table Grid'
        cells = [table.rows[i].cells for i in range(4)]
        cells[0][0].text, cells[0][1].text = "Chỉ số", "Giá trị"
        cells[1][0].text, cells[1][1].text = "Tổng số mẫu (N)", str(total_rows)
        cells[2][0].text, cells[2][1].text = "Tổng số biến số (M)", str(total_cols)
        cells[3][0].text, cells[3][1].text = "Tình trạng", "Dữ liệu đã được làm sạch và kiểm tra logic"
        for i in range(2): set_cell_background(table.rows[0].cells[i], "D9D9D9")
        
        doc.add_heading('3.2. Nhận xét về phương pháp nghiên cứu', level=2)
        add_ai_paragraphs(doc, analysis_meth)

        if stats_type and stats_data:
            doc.add_heading('3.3. Kết quả các kiểm định SPSS nâng cao', level=2)
            if stats_type == 'reliability':
                doc.add_paragraph(f"Hệ số Cronbach's Alpha: {stats_data.get('alpha')} ({stats_data.get('status')})").bold = True
            elif stats_type == 'regression':
                doc.add_paragraph(f"Mô hình Hồi quy (R2={stats_data.get('r_squared')})").bold = True
                t = doc.add_table(rows=1, cols=5); t.style = 'Table Grid'
                h = t.rows[0].cells; h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = "Biến", "Beta", "S.E", "P", "Sig"
                for c in h: set_cell_background(c, "E7E6E6")
                for c in stats_data.get('coefficients', []):
                    r = t.add_row().cells; r[0].text, r[1].text, r[2].text, r[3].text, r[4].text = str(c['variable']), str(c['coefficient']), str(c['std_err']), str(c['p_value']), str(c['sig'])
                    if float(str(c['sig']).replace('<','')) < 0.05:
                        for cell in r: 
                            for p in cell.paragraphs: 
                                for run in p.runs: run.bold = True

        # IV.
        add_heading(doc, 'PHÂN TÍCH CHUYÊN SÂU & BIỆN LUẬN', 1, "IV.")
        add_ai_paragraphs(doc, analysis_iv)

        # V.
        add_heading(doc, 'KẾT LUẬN & KIẾN NGHỊ', 1, "V.")
        add_ai_paragraphs(doc, analysis_v)

        # SAVE
        f = io.BytesIO(); doc.save(f); f.seek(0)
        return send_file(f, as_attachment=True, download_name=f'Bao_cao_NCKH_HOU_Full.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        return {"error": str(e)}, 500